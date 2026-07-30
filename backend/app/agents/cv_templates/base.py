"""CV DOCX renderer — template-aware, parameterised by TemplateConfig."""

import io
import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schemas.tailored import TailoredResume

_SECTION_HEADERS = {
    "CONTACT", "CONTACT INFORMATION", "CONTACT DETAILS",
    "SUMMARY", "PROFESSIONAL SUMMARY", "EXECUTIVE SUMMARY", "OBJECTIVE", "CAREER OBJECTIVE",
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY",
    "EDUCATION", "ACADEMIC BACKGROUND", "ACADEMIC QUALIFICATIONS",
    "SKILLS", "CORE SKILLS", "TECHNICAL SKILLS", "KEY SKILLS",
    "CORE SKILLS & TECHNOLOGIES", "CORE COMPETENCIES", "SKILLS & EXPERTISE",
    "CERTIFICATIONS", "LICENSES & CERTIFICATIONS", "PROFESSIONAL CERTIFICATIONS",
    "ACHIEVEMENTS", "KEY ACHIEVEMENTS", "ACCOMPLISHMENTS", "AWARDS",
    "PROJECTS", "KEY PROJECTS", "NOTABLE PROJECTS", "SELECTED PROJECTS",
    "LANGUAGES", "INTERESTS", "HOBBIES", "REFERENCES",
    "LEADERSHIP", "LEADERSHIP & MANAGEMENT", "VOLUNTEER", "VOLUNTEER EXPERIENCE",
    "PUBLICATIONS", "SPEAKING ENGAGEMENTS", "PROFESSIONAL AFFILIATIONS",
}

_SOURCE_CITATION_RE = re.compile(r'\.\s*Source:\s*[^.]+$', re.IGNORECASE)


@dataclass
class TemplateConfig:
    id: str
    name: str
    font_name: str = "Calibri"
    # Colored header band (None = plain white header)
    header_bg: Optional[tuple] = None
    header_text_color: tuple = (255, 255, 255)
    # Section heading styling
    heading_color: tuple = (0, 0, 0)
    heading_style: str = "rule"  # "rule" | "color_rule" | "left_border" | "shaded_box"
    rule_color: Optional[tuple] = None  # None = heading_color
    # Name block
    name_size: float = 22.0
    name_align: str = "center"   # "center" | "left"
    designation_color: Optional[tuple] = None  # None = heading_color
    contact_align: str = "center"
    # Two-column layout: (sidebar_width_fraction, sidebar_bg_color)
    two_column: Optional[tuple] = None
    # Header bottom rule (for plain white headers that have a border-bottom)
    header_rule_color: Optional[tuple] = None  # None = no rule
    header_rule_sz: int = 16  # border size in 8ths of a point (16 = 2pt)
    # Misc
    heading_uppercase: bool = True
    serif_font: bool = False


class CVDocxRenderer:
    """Renders a TailoredResume as a DOCX using the given TemplateConfig."""

    def __init__(self, config: TemplateConfig) -> None:
        self.cfg = config

    # ── Public entry point ────────────────────────────────────────────

    def render(
        self,
        tailored: TailoredResume,
        candidate_name: str,
        contact_info: Optional[dict] = None,
    ) -> bytes:
        doc = self._base_doc()
        cfg = self.cfg
        contact_info = contact_info or {}

        if cfg.two_column:
            self._render_two_column(doc, tailored, candidate_name, contact_info)
        else:
            self._render_single_column(doc, tailored, candidate_name, contact_info)

        return self._to_bytes(doc)

    # ── Single-column layout ──────────────────────────────────────────

    def _render_single_column(self, doc, tailored, candidate_name, contact_info):
        cfg = self.cfg

        if cfg.header_bg:
            self._add_colored_header(doc, candidate_name, contact_info)
        else:
            self._add_plain_header(doc, candidate_name, contact_info)

        full_text = tailored.full_text or self._sections_to_text(tailored)
        self._render_body(doc, full_text, candidate_name, skip_contact=bool(cfg.header_bg))

    # ── Two-column layout ─────────────────────────────────────────────

    def _render_two_column(self, doc, tailored, candidate_name, contact_info):
        cfg = self.cfg
        sidebar_frac, sidebar_bg = cfg.two_column

        page_width = 6.5  # usable inches after 1" margins
        sidebar_w = Inches(page_width * sidebar_frac)
        main_w = Inches(page_width * (1 - sidebar_frac))

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _remove_table_borders(table)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        left_cell.width = sidebar_w
        right_cell.width = main_w
        _set_cell_bg(left_cell, sidebar_bg)

        # Sidebar: name + designation + skills/certs
        self._fill_sidebar(left_cell, candidate_name, contact_info, tailored)

        # Main: summary + experience + education
        self._fill_main_col(right_cell, tailored, candidate_name)

    def _fill_sidebar(self, cell, candidate_name, contact_info, tailored):
        cfg = self.cfg
        _, sidebar_bg = cfg.two_column
        text_color = (240, 245, 255)
        accent_color = (125, 211, 252)  # light blue

        p = cell.paragraphs[0]
        _clear_para(p)
        run = p.add_run(candidate_name)
        run.font.name = cfg.font_name
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

        designation = contact_info.get("designation", "")
        if designation:
            p2 = cell.add_paragraph()
            r = p2.add_run(designation)
            r.font.name = cfg.font_name
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*accent_color)

        # Extract skills/certs sections from full_text
        full_text = tailored.full_text or ""
        sidebar_sections = {"SKILLS", "CORE SKILLS", "TECHNICAL SKILLS", "KEY SKILLS",
                            "CORE COMPETENCIES", "CERTIFICATIONS", "LANGUAGES"}
        in_sidebar = False
        for line in full_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            upper = stripped.upper()
            is_hdr = (stripped.isupper() and 2 < len(stripped) <= 70) or upper in _SECTION_HEADERS
            if is_hdr:
                in_sidebar = upper in sidebar_sections
                if in_sidebar:
                    sh = cell.add_paragraph()
                    _clear_para(sh)
                    sr = sh.add_run(upper)
                    sr.font.name = cfg.font_name
                    sr.font.size = Pt(9)
                    sr.bold = True
                    sr.font.color.rgb = RGBColor(*accent_color)
            elif in_sidebar:
                sp = cell.add_paragraph()
                _clear_para(sp)
                sr2 = sp.add_run(stripped)
                sr2.font.name = cfg.font_name
                sr2.font.size = Pt(9)
                sr2.font.color.rgb = RGBColor(*text_color)

    def _fill_main_col(self, cell, tailored, candidate_name):
        cfg = self.cfg
        exclude = {"SKILLS", "CORE SKILLS", "TECHNICAL SKILLS", "KEY SKILLS",
                   "CORE COMPETENCIES", "CERTIFICATIONS", "LANGUAGES",
                   "CONTACT", "CONTACT INFORMATION", "CONTACT DETAILS"}

        full_text = tailored.full_text or self._sections_to_text(tailored)
        in_excluded = False
        first_para = True

        for line in full_text.split("\n"):
            stripped = _SOURCE_CITATION_RE.sub(".", line.strip()).strip()
            if not stripped or stripped == ".":
                continue
            if stripped.upper() == candidate_name.upper():
                continue

            upper = stripped.upper()
            is_hdr = (stripped.isupper() and 2 < len(stripped) <= 70) or upper in _SECTION_HEADERS
            if is_hdr:
                in_excluded = upper in exclude
                if not in_excluded:
                    if first_para:
                        p = cell.paragraphs[0]
                        _clear_para(p)
                        first_para = False
                    else:
                        p = cell.add_paragraph()
                    self._apply_heading_style(p, stripped)
            elif not in_excluded:
                if first_para:
                    p = cell.paragraphs[0]
                    _clear_para(p)
                    first_para = False
                else:
                    p = cell.add_paragraph()
                r = p.add_run(stripped)
                r.font.name = cfg.font_name
                r.font.size = Pt(10)

    # ── Header helpers ────────────────────────────────────────────────

    def _add_colored_header(self, doc, candidate_name, contact_info):
        cfg = self.cfg
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        _remove_table_borders(table)
        cell = table.cell(0, 0)
        _set_cell_bg(cell, cfg.header_bg)
        cell.top_margin = Pt(12)
        cell.bottom_margin = Pt(12)
        cell.left_margin = Pt(16)

        p = cell.paragraphs[0]
        _clear_para(p)
        r = p.add_run(candidate_name.upper() if cfg.heading_uppercase else candidate_name)
        r.font.name = cfg.font_name
        r.font.size = Pt(cfg.name_size)
        r.bold = True
        r.font.color.rgb = RGBColor(*cfg.header_text_color)

        designation = contact_info.get("designation", "")
        if designation:
            p2 = cell.add_paragraph()
            dr = p2.add_run(designation)
            dr.font.name = cfg.font_name
            dr.font.size = Pt(12)
            desig_color = cfg.designation_color or (200, 220, 255)
            dr.font.color.rgb = RGBColor(*desig_color)

        contact_line = _build_contact_line(contact_info)
        if contact_line:
            p3 = cell.add_paragraph()
            cr = p3.add_run(contact_line)
            cr.font.name = cfg.font_name
            cr.font.size = Pt(10)
            cr.font.color.rgb = RGBColor(180, 200, 220)

        doc.add_paragraph()  # spacer after header table

    def _add_plain_header(self, doc, candidate_name, contact_info):
        cfg = self.cfg
        align = WD_ALIGN_PARAGRAPH.CENTER if cfg.name_align == "center" else WD_ALIGN_PARAGRAPH.LEFT

        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(candidate_name.upper() if cfg.heading_uppercase else candidate_name)
        r.font.name = cfg.font_name
        r.font.size = Pt(cfg.name_size)
        r.bold = cfg.name_align == "center"
        r.font.color.rgb = RGBColor(*cfg.heading_color)

        designation = contact_info.get("designation", "")
        if designation:
            p2 = doc.add_paragraph()
            p2.alignment = align
            dr = p2.add_run(designation)
            dr.font.name = cfg.font_name
            dr.font.size = Pt(12)
            dr.bold = True
            desig_color = cfg.designation_color or cfg.heading_color
            dr.font.color.rgb = RGBColor(*desig_color)

        contact_line = _build_contact_line(contact_info)
        if contact_line:
            p3 = doc.add_paragraph()
            p3.alignment = align
            cr = p3.add_run(contact_line)
            cr.font.name = cfg.font_name
            cr.font.size = Pt(10)
            cr.font.color.rgb = RGBColor(90, 90, 90)

        # Bottom border rule under the header block (e.g. ATS Classic, Harvard)
        if cfg.header_rule_color:
            rule_p = doc.add_paragraph()
            rule_p.alignment = align
            _add_bottom_border(rule_p, cfg.header_rule_color, cfg.header_rule_sz)

    # ── Body renderer ─────────────────────────────────────────────────

    def _render_body(self, doc, full_text, candidate_name, skip_contact=False):
        cfg = self.cfg
        candidate_upper = candidate_name.upper().replace(" ", "")
        in_contact = False      # inside CONTACT section (skip when skip_contact=True)
        in_header_sec = False   # inside HEADER section (always skip — already in doc header)

        for line in full_text.split("\n"):
            stripped = _SOURCE_CITATION_RE.sub(".", line.strip()).strip()
            if not stripped or stripped == ".":
                continue
            if stripped.upper().replace(" ", "") == candidate_upper:
                continue

            upper = stripped.upper()
            is_hdr = (stripped.isupper() and 2 < len(stripped) <= 70) or upper in _SECTION_HEADERS

            if is_hdr:
                if upper == "HEADER":
                    in_header_sec = True
                    in_contact = False
                    continue
                if upper in {"CONTACT", "CONTACT INFORMATION", "CONTACT DETAILS"}:
                    in_contact = True
                    in_header_sec = False
                    continue
                in_contact = False
                in_header_sec = False
                doc.add_paragraph()  # spacer
                p = doc.add_paragraph()
                self._apply_heading_style(p, stripped)
            elif in_header_sec:
                continue  # always skip content inside HEADER section
            elif in_contact and skip_contact:
                continue
            else:
                p = doc.add_paragraph()
                r = p.add_run(stripped)
                r.font.name = cfg.font_name
                r.font.size = Pt(11)

    def _apply_heading_style(self, para, heading_text):
        cfg = self.cfg
        text = heading_text.upper() if cfg.heading_uppercase else heading_text
        r = para.add_run(text)
        r.font.name = cfg.font_name
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(*cfg.heading_color)

        if cfg.heading_style == "shaded_box":
            _shade_paragraph(para, cfg.heading_color)
            r.font.color.rgb = RGBColor(255, 255, 255)
        elif cfg.heading_style == "left_border":
            border_color = cfg.rule_color or cfg.heading_color
            _add_left_border(para, border_color)
        elif cfg.heading_style in ("rule", "color_rule"):
            border_color = cfg.rule_color or cfg.heading_color
            sz = 16 if cfg.heading_style == "color_rule" else 8
            _add_bottom_border(para, border_color, sz)

    # ── Utilities ─────────────────────────────────────────────────────

    @staticmethod
    def _sections_to_text(tailored: TailoredResume) -> str:
        parts = []
        for s in tailored.sections:
            if s.section_name:
                parts.append(s.section_name.upper())
            if s.content:
                parts.append(s.content)
        return "\n\n".join(parts)

    def _base_doc(self) -> Document:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Inches(1.0)
            sec.bottom_margin = Inches(1.0)
            sec.left_margin = Inches(1.0)
            sec.right_margin = Inches(1.0)
        style = doc.styles["Normal"]
        style.font.name = self.cfg.font_name
        style.font.size = Pt(11)
        return doc

    @staticmethod
    def _to_bytes(doc: Document) -> bytes:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


# ── XML helpers ───────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: tuple) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _shade_paragraph(para, rgb: tuple) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_left_border(para, rgb: tuple) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_bottom_border(para, rgb: tuple, sz: int = 8) -> None:
    """Add a bottom border to a paragraph (rule/color_rule heading styles)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)


def _clear_para(para) -> None:
    for child in list(para._p):
        para._p.remove(child)


def _build_contact_line(contact_info: dict) -> str:
    parts = []
    for key in ("location", "email", "phone", "linkedin"):
        val = contact_info.get(key, "")
        if val:
            parts.append(val)
    return "  ·  ".join(parts)
