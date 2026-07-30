"""Package generation agent — creates DOCX/PDF documents from analysis results.

Does NOT call an LLM. Generates ATS-safe documents using python-docx with:
- Single column layout
- Calibri 11pt font
- Standard section headers
- No tables or complex formatting
"""

import io
import logging
import zipfile
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel

from app.schemas.cover_letter import CoverLetter
from app.schemas.interview import InterviewGuide, InterviewQuestion
from app.schemas.tailored import TailoredResume

logger = logging.getLogger(__name__)

# ATS-safe formatting constants
_FONT_NAME = "Calibri"
_FONT_SIZE_BODY = Pt(11)
_FONT_SIZE_HEADING = Pt(14)
_FONT_SIZE_SUBHEADING = Pt(12)
_LINE_SPACING = 1.15
_MARGIN_INCHES = 1.0

# Known ATS resume section header names (uppercase comparison)
_SECTION_HEADERS = {
    "CONTACT", "CONTACT INFORMATION", "CONTACT DETAILS",
    "SUMMARY", "PROFESSIONAL SUMMARY", "EXECUTIVE SUMMARY", "OBJECTIVE", "CAREER OBJECTIVE",
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY",
    "EDUCATION", "ACADEMIC BACKGROUND", "ACADEMIC QUALIFICATIONS",
    "SKILLS", "CORE SKILLS", "TECHNICAL SKILLS", "KEY SKILLS",
    "CORE SKILLS & TECHNOLOGIES", "CORE COMPETENCIES", "SKILLS & EXPERTISE",
    "CERTIFICATIONS", "LICENSES & CERTIFICATIONS", "PROFESSIONAL CERTIFICATIONS",
    "ACHIEVEMENTS", "KEY ACHIEVEMENTS", "ACCOMPLISHMENTS", "AWARDS",
    "PROJECTS", "KEY PROJECTS", "NOTABLE PROJECTS", "SELECTED PROJECTS",
    "LANGUAGES", "INTERESTS", "HOBBIES", "REFERENCES",
    "LEADERSHIP", "LEADERSHIP & MANAGEMENT", "VOLUNTEER", "VOLUNTEER EXPERIENCE",
    "PUBLICATIONS", "SPEAKING ENGAGEMENTS", "PROFESSIONAL AFFILIATIONS",
}


class PackageInput(BaseModel):
    """Input schema for the package generation agent."""

    tailored_resume: TailoredResume
    cover_letter: CoverLetter
    interview_guide: InterviewGuide
    candidate_name: Optional[str] = None


class PackageOutput(BaseModel):
    """Output schema for the package generation agent."""

    resume_docx: bytes
    cover_letter_docx: bytes
    interview_guide_docx: bytes
    zip_bundle: bytes


class PackageGenerationAgent:
    """Generates ATS-safe DOCX documents and a ZIP bundle.

    Does NOT call an LLM. Uses python-docx for document generation
    with ATS-safe formatting: single column, Calibri 11pt, standard
    section headers, no tables or complex formatting.
    """

    agent_name = "package_generation"

    def __init__(self) -> None:
        """Initialize the package generation agent (no LLM needed)."""
        self._total_input_tokens = 0
        self._total_output_tokens = 0

    @property
    def total_input_tokens(self) -> int:
        """Always 0 — this agent does not call an LLM."""
        return 0

    @property
    def total_output_tokens(self) -> int:
        """Always 0 — this agent does not call an LLM."""
        return 0

    async def execute(self, input_data: PackageInput) -> PackageOutput:
        """Generate all documents and bundle them into a ZIP.

        Args:
            input_data: Tailored resume, cover letter, and interview guide data.

        Returns:
            PackageOutput with bytes for each document and a ZIP bundle.
        """
        candidate = input_data.candidate_name or "Candidate"

        resume_docx = self._generate_resume_docx(input_data.tailored_resume, candidate)
        cover_letter_docx = self._generate_cover_letter_docx(input_data.cover_letter, candidate)
        interview_guide_docx = self._generate_interview_guide_docx(
            input_data.interview_guide, candidate
        )

        zip_bundle = self._create_zip_bundle(
            resume_docx=resume_docx,
            cover_letter_docx=cover_letter_docx,
            interview_guide_docx=interview_guide_docx,
            candidate_name=candidate,
        )

        logger.info("Generated document package for %s", candidate)

        return PackageOutput(
            resume_docx=resume_docx,
            cover_letter_docx=cover_letter_docx,
            interview_guide_docx=interview_guide_docx,
            zip_bundle=zip_bundle,
        )

    def _generate_resume_docx(self, tailored: TailoredResume, candidate: str) -> bytes:
        """Generate an ATS-safe resume DOCX document.

        Args:
            tailored: The tailored resume data.
            candidate: Candidate name for the document.

        Returns:
            DOCX file content as bytes.
        """
        doc = self._create_base_document()

        # Add candidate name as title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(candidate)
        run.font.size = _FONT_SIZE_HEADING
        run.font.name = _FONT_NAME
        run.bold = True

        # Sections are well-formed only when they have BOTH named headings AND body content.
        # The LLM sometimes provides sections with content but no section_name, or
        # section names in the content field with no body text.
        sections_have_names = any(s.section_name.strip() for s in tailored.sections)
        sections_have_content = any(
            len(s.content.strip()) > 60 or "\n" in s.content
            for s in tailored.sections
            if s.content
        )

        if sections_have_names and sections_have_content:
            for section in tailored.sections:
                if section.section_name.strip():
                    heading = doc.add_paragraph()
                    heading_run = heading.add_run(section.section_name.upper())
                    heading_run.font.size = _FONT_SIZE_SUBHEADING
                    heading_run.font.name = _FONT_NAME
                    heading_run.bold = True
                    separator = doc.add_paragraph()
                    separator_run = separator.add_run("─" * 60)
                    separator_run.font.size = Pt(8)
                    separator_run.font.color.rgb = RGBColor(180, 180, 180)

                for para_text in section.content.split("\n"):
                    if para_text.strip():
                        para = doc.add_paragraph()
                        para_run = para.add_run(para_text.strip())
                        para_run.font.size = _FONT_SIZE_BODY
                        para_run.font.name = _FONT_NAME
        elif tailored.full_text:
            # Fallback: render the full_text field, auto-detecting section headers
            logger.info("Sections have no body content — falling back to full_text rendering")
            self._render_full_text_as_resume(doc, tailored.full_text, candidate)
        else:
            para = doc.add_paragraph()
            para.add_run("Resume content could not be generated.")

        return self._doc_to_bytes(doc)

    def _render_full_text_as_resume(
        self, doc: Document, full_text: str, candidate: str
    ) -> None:
        """Render full_text as ATS-safe sections by detecting header lines.

        Header block (CONTACT section) gets special centred formatting:
        - designation line → bold Calibri 12pt centred
        - contact info lines → grey Calibri 10pt centred

        Args:
            doc: Document to write into.
            full_text: The complete resume text.
            candidate: Candidate name (skipped wherever it appears — already the title).
        """
        import re
        candidate_upper = candidate.upper().replace(" ", "")
        # Strip trailing ". Source: …" citations the claim-verification agent embeds
        _source_citation_re = re.compile(r'\.\s*Source:\s*[^.]+$', re.IGNORECASE)

        in_contact_section = False
        header_body_count = 0   # lines rendered inside the CONTACT block (0=designation)

        for line in full_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            stripped = _source_citation_re.sub(".", stripped).strip()
            if not stripped or stripped == ".":
                continue

            # Skip the candidate name wherever it appears — already the doc title
            if (
                stripped.upper() == candidate.upper()
                or stripped.upper().replace(" ", "") == candidate_upper
            ):
                continue

            upper = stripped.upper()
            is_header = (
                stripped.isupper()
                and 2 < len(stripped) <= 70
                and not stripped.startswith(("•", "-", "*", "·"))
            ) or upper in _SECTION_HEADERS

            if is_header:
                # Detect start of CONTACT block — skip the visible header for CONTACT
                # (the name is already the document title)
                if upper in {"CONTACT", "CONTACT INFORMATION", "CONTACT DETAILS"}:
                    in_contact_section = True
                    header_body_count = 0
                    continue  # Don't render "CONTACT" as a visible heading

                in_contact_section = False

                # Add a spacer before non-contact section headings
                doc.add_paragraph()

                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                heading_run = heading.add_run(stripped.upper())
                heading_run.font.size = _FONT_SIZE_SUBHEADING
                heading_run.font.name = _FONT_NAME
                heading_run.bold = True
                heading_run.font.color.rgb = RGBColor(31, 56, 100)   # dark navy

                sep = doc.add_paragraph()
                sep_run = sep.add_run("─" * 60)
                sep_run.font.size = Pt(8)
                sep_run.font.color.rgb = RGBColor(31, 56, 100)
            elif in_contact_section:
                if header_body_count == 0:
                    # First body line = designation / current job title
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(stripped)
                    run.font.name = _FONT_NAME
                    run.font.size = _FONT_SIZE_SUBHEADING   # 12pt
                    run.bold = True
                    run.font.color.rgb = RGBColor(31, 56, 100)
                else:
                    # Subsequent lines = contact info (email, phone, LinkedIn…)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(stripped)
                    run.font.name = _FONT_NAME
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(90, 90, 90)
                header_body_count += 1
            else:
                para = doc.add_paragraph()
                para_run = para.add_run(stripped)
                para_run.font.size = _FONT_SIZE_BODY
                para_run.font.name = _FONT_NAME

    def _generate_cover_letter_docx(self, cover_letter: CoverLetter, candidate: str) -> bytes:
        """Generate a professionally formatted cover letter DOCX.

        Args:
            cover_letter: The cover letter data.
            candidate: Candidate name for the document.

        Returns:
            DOCX file content as bytes.
        """
        import datetime
        doc = self._create_base_document()

        def _add_text(text, bold=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
            p = doc.add_paragraph()
            p.alignment = align
            r = p.add_run(text)
            r.font.name = _FONT_NAME
            r.font.size = size or _FONT_SIZE_BODY
            r.bold = bold
            return p

        # ── Sender block (From) ──────────────────────────────────────
        _add_text(candidate, bold=True, size=_FONT_SIZE_SUBHEADING)
        _add_text("Chennai, India  |  suresh.natarajan19@gmail.com  |  linkedin.com/in/sureshnsm")
        doc.add_paragraph()

        # ── Date ─────────────────────────────────────────────────────
        date_str = datetime.date.today().strftime("%B %d, %Y")
        _add_text(date_str)
        doc.add_paragraph()

        # ── Recipient block (To) ─────────────────────────────────────
        company = cover_letter.company_name or "Hiring Company"
        role = cover_letter.role_title or "Open Position"
        _add_text("Hiring Manager", bold=True)
        _add_text(company)
        doc.add_paragraph()

        # ── Subject ──────────────────────────────────────────────────
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"Re: Application for {role}")
        subject_run.font.name = _FONT_NAME
        subject_run.font.size = _FONT_SIZE_BODY
        subject_run.bold = True
        subject_run.underline = True
        doc.add_paragraph()

        # ── Salutation ───────────────────────────────────────────────
        _add_text("Dear Hiring Manager,")
        doc.add_paragraph()

        # ── Body ─────────────────────────────────────────────────────
        body = cover_letter.content.strip()
        if body:
            for para_text in body.split("\n\n"):
                if para_text.strip():
                    _add_text(para_text.strip())
                    doc.add_paragraph()
        else:
            _add_text("[Cover letter content not available — please re-run the analysis.]")

        # ── Closing ──────────────────────────────────────────────────
        _add_text("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph()
        _add_text(candidate, bold=True)
        _add_text("Agentic AI Architect | Solution Architect")

        return self._doc_to_bytes(doc)

    def _generate_interview_guide_docx(
        self, guide: InterviewGuide, candidate: str
    ) -> bytes:
        """Generate an interview preparation guide DOCX document.

        Args:
            guide: The interview guide data.
            candidate: Candidate name for the document.

        Returns:
            DOCX file content as bytes.
        """
        doc = self._create_base_document()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Interview Preparation Guide")
        run.font.size = _FONT_SIZE_HEADING
        run.font.name = _FONT_NAME
        run.bold = True

        doc.add_paragraph()  # Spacer

        # Preparation tips
        tips_heading = doc.add_paragraph()
        tips_run = tips_heading.add_run("PREPARATION TIPS")
        tips_run.font.size = _FONT_SIZE_SUBHEADING
        tips_run.font.name = _FONT_NAME
        tips_run.bold = True

        for tip in guide.preparation_tips:
            para = doc.add_paragraph()
            para_run = para.add_run(f"• {tip}")
            para_run.font.size = _FONT_SIZE_BODY
            para_run.font.name = _FONT_NAME

        doc.add_paragraph()  # Spacer

        # Behavioral questions
        self._add_question_section(doc, "BEHAVIORAL QUESTIONS", guide.behavioral_questions)

        doc.add_paragraph()  # Spacer

        # Technical questions
        self._add_question_section(doc, "TECHNICAL QUESTIONS", guide.technical_questions)

        return self._doc_to_bytes(doc)

    def _add_question_section(
        self, doc: Document, heading_text: str, questions: list[InterviewQuestion]
    ) -> None:
        """Add a section of interview questions to the document.

        Args:
            doc: The Document to add to.
            heading_text: Section heading text.
            questions: List of interview questions.
        """
        heading = doc.add_paragraph()
        heading_run = heading.add_run(heading_text)
        heading_run.font.size = _FONT_SIZE_SUBHEADING
        heading_run.font.name = _FONT_NAME
        heading_run.bold = True

        for i, q in enumerate(questions, 1):
            # Question text
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{i}. {q.question}")
            q_run.font.size = _FONT_SIZE_BODY
            q_run.font.name = _FONT_NAME
            q_run.bold = True

            # Source info
            if q.resume_evidence:
                evidence_para = doc.add_paragraph()
                evidence_run = evidence_para.add_run(f"   Evidence: {q.resume_evidence}")
                evidence_run.font.size = Pt(10)
                evidence_run.font.name = _FONT_NAME
                evidence_run.font.color.rgb = RGBColor(80, 80, 80)

            # STAR skeleton
            if q.star_skeleton:
                for key in ("situation", "task", "action", "result"):
                    val = q.star_skeleton.get(key, "")
                    if val:
                        star_para = doc.add_paragraph()
                        star_run = star_para.add_run(f"   {key.capitalize()}: {val}")
                        star_run.font.size = Pt(10)
                        star_run.font.name = _FONT_NAME

            # Note
            if q.note:
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(f"   Note: {q.note}")
                note_run.font.size = Pt(10)
                note_run.font.name = _FONT_NAME
                note_run.italic = True

    def _create_base_document(self) -> Document:
        """Create a base DOCX document with ATS-safe formatting.

        Returns:
            Document with standard margins and formatting.
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(_MARGIN_INCHES)
            section.bottom_margin = Inches(_MARGIN_INCHES)
            section.left_margin = Inches(_MARGIN_INCHES)
            section.right_margin = Inches(_MARGIN_INCHES)

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = _FONT_NAME
        font.size = _FONT_SIZE_BODY

        # Set line spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = _LINE_SPACING

        return doc

    def _doc_to_bytes(self, doc: Document) -> bytes:
        """Serialize a Document to bytes.

        Args:
            doc: python-docx Document.

        Returns:
            DOCX file content as bytes.
        """
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _create_zip_bundle(
        self,
        resume_docx: bytes,
        cover_letter_docx: bytes,
        interview_guide_docx: bytes,
        candidate_name: str,
    ) -> bytes:
        """Bundle all documents into a ZIP file.

        Args:
            resume_docx: Resume document bytes.
            cover_letter_docx: Cover letter document bytes.
            interview_guide_docx: Interview guide document bytes.
            candidate_name: Candidate name for file naming.

        Returns:
            ZIP file content as bytes.
        """
        buffer = io.BytesIO()
        safe_name = candidate_name.replace(" ", "_").lower()

        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{safe_name}_resume_tailored.docx", resume_docx)
            zf.writestr(f"{safe_name}_cover_letter.docx", cover_letter_docx)
            zf.writestr(f"{safe_name}_interview_guide.docx", interview_guide_docx)

        buffer.seek(0)
        return buffer.read()
